import hashlib
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document as DocxDocument
from pypdf import PdfReader

from rag_kit.models import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_path(path: Path) -> List[Document]:
    path = path.resolve()
    if path.is_dir():
        documents: List[Document] = []
        for child in sorted(path.rglob("*")):
            if child.is_file() and child.suffix.lower() in SUPPORTED_EXTENSIONS:
                documents.extend(load_file(child))
        return documents
    return load_file(path)


def load_file(path: Path) -> List[Document]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported file type: {suffix}")

    if suffix in {".txt", ".md"}:
        return [_document(path, path.read_text(encoding="utf-8", errors="ignore"))]
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return [_document(path, _load_docx_text(path))]
    raise ValueError(f"unsupported file type: {suffix}")


def _load_pdf(path: Path) -> List[Document]:
    reader = PdfReader(str(path))
    documents: List[Document] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                _document(path, text, extra_metadata={"page": page_number})
            )
    return documents


def _load_docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _document(path: Path, text: str, extra_metadata: Optional[Dict] = None) -> Document:
    digest = hashlib.sha1(f"{path}:{extra_metadata}:{text[:200]}".encode("utf-8")).hexdigest()[:16]
    metadata = {"source": str(path)}
    if extra_metadata:
        metadata.update(extra_metadata)
    return Document(id=digest, text=text, metadata=metadata)
